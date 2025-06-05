from flask import Blueprint, request, jsonify, current_app
from flask_jwt_extended import jwt_required, get_jwt_identity, get_jwt
from sqlalchemy import or_, distinct
from datetime import datetime, timedelta
from collections import Counter

from app.models import db, Province, LedgerEntry, User

meta_bp = Blueprint('meta', __name__)

@meta_bp.route('/provinces', methods=['GET'])
@jwt_required()
def get_provinces():
    """获取省份列表"""
    try:
        # 从数据库获取省份
        provinces = Province.query.order_by(Province.name).all()
        province_list = [p.name for p in provinces]
        
        # 如果数据库为空，使用配置中的默认列表
        if not province_list:
            province_list = current_app.config.get('PROVINCES', [])
        
        return jsonify({
            'code': 0,
            'message': 'success',
            'data': province_list
        }), 200
        
    except Exception as e:
        current_app.logger.error(f'Get provinces error: {str(e)}')
        return jsonify({
            'code': 500,
            'message': 'Failed to get provinces',
            'data': None
        }), 500

@meta_bp.route('/suggestions/project_items', methods=['GET'])
@jwt_required()
def get_project_suggestions():
    """获取项目名称搜索建议，支持省份过滤和频率排序"""
    try:
        user_id = get_jwt_identity()
        user_role = get_jwt().get('role', 'user')
        query_text = request.args.get('query', '').strip()
        province = request.args.get('province', '').strip()
        if not query_text:
            return jsonify({'code': 0, 'message': 'success', 'data': []}), 200
        search_pattern = f'%{query_text}%'
        query = LedgerEntry.query
        if user_role == 'user':
            query = query.filter_by(user_id=user_id)
        if province:
            query = query.filter(LedgerEntry.province == province)
        # 统计建议项出现频率
        suggestions = query.filter(
            or_(
                LedgerEntry.project_name.ilike(search_pattern),
                LedgerEntry.location.ilike(search_pattern),
                LedgerEntry.specific_matters.ilike(search_pattern)
            )
        ).all()
        counter = Counter()
        for entry in suggestions:
            if entry.project_name and query_text.lower() in entry.project_name.lower():
                counter[entry.project_name] += 1
            if entry.location and query_text.lower() in entry.location.lower():
                counter[entry.location] += 1
            if entry.specific_matters and query_text.lower() in entry.specific_matters.lower():
                sentences = entry.specific_matters.split('。')
                for sentence in sentences:
                    if query_text.lower() in sentence.lower() and len(sentence) < 100:
                        counter[sentence.strip()] += 1
        formatted_suggestions = [
            {"value": value, "count": count}
            for value, count in counter.most_common(20)
        ]
        return jsonify({'code': 0, 'message': 'success', 'data': formatted_suggestions}), 200
    except Exception as e:
        current_app.logger.error(f'Get project suggestions error: {str(e)}')
        return jsonify({'code': 500, 'message': 'Failed to get suggestions', 'data': None}), 500

@meta_bp.route('/nature-options', methods=['GET'])
@jwt_required()
def get_nature_options():
    """获取性质选项列表"""
    try:
        # 预定义的性质选项
        nature_options = [
            '会议纪要',
            '工作安排',
            '问题反馈',
            '质量管理',
            '临时任务',
            '研讨交流',
            '资料交接',
            '人员对接',
            '常规项目工作',
            '其他'
        ]
        
        return jsonify({
            'code': 0,
            'message': 'success',
            'data': nature_options
        }), 200
        
    except Exception as e:
        current_app.logger.error(f'Get nature options error: {str(e)}')
        return jsonify({
            'code': 500,
            'message': 'Failed to get nature options',
            'data': None
        }), 500

@meta_bp.route('/export/ledger', methods=['GET'])
@jwt_required()
def export_ledger():
    """导出台账数据（支持CSV、Excel、Word格式，支持搜索过滤）"""
    try:
        import csv
        import io
        from flask import Response
        from openpyxl import Workbook
        from docx import Document
        from docx.shared import Inches
        user_id = get_jwt_identity()
        user_role = get_jwt().get('role', 'user')
        if user_role not in ['admin', 'power_user', 'user']:
            return jsonify({'code': 403, 'message': 'Insufficient permissions', 'data': None}), 403
        # 获取搜索参数
        project_name = request.args.get('project_name', '').strip()
        location = request.args.get('location', '').strip()
        province = request.args.get('province', '').strip()
        recorder = request.args.get('recorder', '').strip()
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')
        # 构建查询
        query = LedgerEntry.query
        if user_role == 'user':
            query = query.filter_by(user_id=user_id)
        if project_name:
            project_pattern = f'%{project_name}%'
            query = query.filter(LedgerEntry.project_name.ilike(project_pattern))
        if location:
            location_pattern = f'%{location}%'
            query = query.filter(LedgerEntry.location.ilike(location_pattern))
        if province:
            query = query.filter_by(province=province)
        if recorder:
            query = query.join(User).filter(User.username == recorder)
        if start_date:
            try:
                start = datetime.strptime(start_date, '%Y-%m-%d').date()
                query = query.filter(LedgerEntry.date >= start)
            except ValueError:
                pass
        if end_date:
            try:
                end = datetime.strptime(end_date, '%Y-%m-%d').date()
                query = query.filter(LedgerEntry.date <= end)
            except ValueError:
                pass
        entries = query.order_by(LedgerEntry.created_at.desc()).all()
        current_app.logger.info(f'导出参数: project_name={project_name}, location={location}, province={province}, recorder={recorder}, start_date={start_date}, end_date={end_date}, role={user_role}, user_id={user_id}')
        current_app.logger.info(f'导出结果数量: {len(entries)}')
        # 导出格式
        export_format = request.args.get('format', 'csv').lower()
        headers = [
            'ID', '录入人员', '省份', '项目名称', '日期', '地点',
            '涉及人员', '性质', '具体事项', '后续要点', '创建时间'
        ]
        rows = [
            [
                entry.id,
                entry.author.username if entry.author else '',
                entry.province,
                entry.project_name,
                entry.date.strftime('%Y-%m-%d') if entry.date else '',
                entry.location,
                entry.personnel,
                entry.nature,
                entry.specific_matters,
                entry.follow_up_points or '',
                entry.created_at.strftime('%Y-%m-%d %H:%M:%S') if entry.created_at else ''
            ]
            for entry in entries
        ]
        if export_format == 'excel':
            wb = Workbook()
            ws = wb.active
            ws.append(headers)
            # 加粗表头
            from openpyxl.styles import Font, Alignment
            for cell in ws[1]:
                cell.font = Font(bold=True)
                cell.alignment = Alignment(horizontal='center', vertical='center')
            for row in rows:
                ws.append(row)
            # 自动列宽
            for col in ws.columns:
                max_length = 0
                col_letter = col[0].column_letter
                for cell in col:
                    try:
                        max_length = max(max_length, len(str(cell.value)))
                    except:
                        pass
                ws.column_dimensions[col_letter].width = max_length + 2
            output = io.BytesIO()
            wb.save(output)
            output.seek(0)
            return Response(
                output.getvalue(),
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                headers={
                    'Content-Disposition': 'attachment; filename=ledger_export.xlsx',
                    'Content-Type': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
                }
            )
        elif export_format == 'word':
            doc = Document()
            doc.add_heading('台账导出', 0)
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            from docx.shared import Pt
            for i, h in enumerate(headers):
                hdr_cells[i].text = h
                run = hdr_cells[i].paragraphs[0].runs[0]
                run.bold = True
                run.font.size = Pt(11)
            for row in rows:
                row_cells = table.add_row().cells
                for i, val in enumerate(row):
                    row_cells[i].text = str(val)
            # 自动调整列宽（简单实现）
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            return Response(
                output.getvalue(),
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                headers={
                    'Content-Disposition': 'attachment; filename=ledger_export.docx',
                    'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                }
            )
        else:  # 默认CSV
            output = io.StringIO()
            writer = csv.writer(output)
            writer.writerow(headers)
            for row in rows:
                writer.writerow(row)
            output.seek(0)
            # 加 BOM 头防止 Excel 乱码
            csv_content = '\ufeff' + output.getvalue()
            if not rows:
                # 若无数据也输出表头
                csv_content = '\ufeff' + ','.join(headers) + '\n'
            return Response(
                csv_content,
                mimetype='text/csv',
                headers={
                    'Content-Disposition': 'attachment; filename=ledger_export.csv',
                    'Content-Type': 'text/csv; charset=utf-8'
                }
            )
    except Exception as e:
        current_app.logger.error(f'Export ledger error: {str(e)}')
        return jsonify({
            'code': 500,
            'message': 'Failed to export ledger',
            'data': None
        }), 500

@meta_bp.route('/suggestions/locations', methods=['GET'])
@jwt_required()
def get_location_suggestions():
    """获取地点搜索建议，支持省份过滤和频率排序"""
    try:
        user_id = get_jwt_identity()
        user_role = get_jwt().get('role', 'user')
        query_text = request.args.get('query', '').strip()
        province = request.args.get('province', '').strip()
        if not query_text:
            return jsonify({'code': 0, 'message': 'success', 'data': []}), 200
        search_pattern = f'%{query_text}%'
        query = LedgerEntry.query
        if user_role == 'user':
            query = query.filter_by(user_id=user_id)
        if province:
            query = query.filter(LedgerEntry.province == province)
        suggestions = query.filter(LedgerEntry.location.ilike(search_pattern)).all()
        counter = Counter()
        for entry in suggestions:
            if entry.location and query_text.lower() in entry.location.lower():
                counter[entry.location] += 1
        formatted_suggestions = [
            {"value": value, "count": count}
            for value, count in counter.most_common(20)
        ]
        return jsonify({'code': 0, 'message': 'success', 'data': formatted_suggestions}), 200
    except Exception as e:
        current_app.logger.error(f'Get location suggestions error: {str(e)}')
        return jsonify({'code': 500, 'message': 'Failed to get suggestions', 'data': None}), 500

@meta_bp.route('/stats/ledger_by_user', methods=['GET'])
@jwt_required()
def get_ledger_stats_by_user():
    """获取各用户台账统计"""
    try:
        days = request.args.get('days', default=30, type=int)
        start_date = datetime.now() - timedelta(days=days)

        # 查询每个用户的台账数量
        stats = db.session.query(
            User.username,
            db.func.count(LedgerEntry.id).label('count')
        ).join(
            LedgerEntry, User.id == LedgerEntry.user_id
        ).filter(
            LedgerEntry.created_at >= start_date
        ).group_by(
            User.username
        ).all()

        # 查询每个用户的月度趋势
        monthly_stats = db.session.query(
            User.username,
            db.func.strftime('%Y-%m', LedgerEntry.created_at).label('month'),
            db.func.count(LedgerEntry.id).label('count')
        ).join(
            LedgerEntry, User.id == LedgerEntry.user_id
        ).filter(
            LedgerEntry.created_at >= start_date
        ).group_by(
            User.username,
            db.func.strftime('%Y-%m', LedgerEntry.created_at)
        ).order_by(
            db.func.strftime('%Y-%m', LedgerEntry.created_at)
        ).all()

        # 处理月度趋势数据
        trend_data = {}
        for username, month, count in monthly_stats:
            if username not in trend_data:
                trend_data[username] = []
            trend_data[username].append({
                'month': month,
                'count': count
            })

        return jsonify({
            'success': True,
            'data': {
                'total': [{'username': username, 'count': count} for username, count in stats],
                'trend': [{'username': username, 'records': records} for username, records in trend_data.items()]
            }
        }), 200

    except Exception as e:
        current_app.logger.error(f'Get ledger stats error: {str(e)}')
        return jsonify({
            'success': False,
            'message': 'Failed to get ledger statistics',
            'error': str(e) if current_app.debug else 'Internal error'
        }), 500