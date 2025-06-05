from flask import Blueprint, request, jsonify, current_app, Response
from flask_jwt_extended import jwt_required, get_jwt_identity
from app.models import LedgerEntry


from app.models import db, User, ActivityLog
from app.utils.decorators import role_required
from app.auth import log_activity
from flask_cors import CORS
import csv
import io
from openpyxl import Workbook
from docx import Document
from docx.shared import Pt

admin_bp = Blueprint('admin', __name__)
CORS(admin_bp)

@admin_bp.route('/logs', methods=['GET'])
@jwt_required()
@role_required(['admin'])
def get_activity_logs():
    """获取活动日志"""
    try:
        # 分页参数
        page = request.args.get('page', 1, type=int)
        limit = request.args.get('limit', 20, type=int)
        
        # 限制每页最大数量
        limit = min(limit, 100)
        
        # 日志级别过滤
        level = request.args.get('level', '').upper()
        
        # 用户过滤
        user_id = request.args.get('user_id', type=int)
        
        # 时间范围
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')
        
        # 构建查询
        query = ActivityLog.query
        
        if level and level in ['INFO', 'WARNING', 'ERROR']:
            query = query.filter_by(level=level)
        
        if user_id:
            query = query.filter_by(user_id=user_id)
        
        if start_date:
            try:
                from datetime import datetime
                start = datetime.strptime(start_date, '%Y-%m-%d')
                query = query.filter(ActivityLog.timestamp >= start)
            except ValueError:
                pass
        
        if end_date:
            try:
                from datetime import datetime
                end = datetime.strptime(end_date, '%Y-%m-%d')
                end = end.replace(hour=23, minute=59, second=59)
                query = query.filter(ActivityLog.timestamp <= end)
            except ValueError:
                pass
        
        # 排序
        query = query.order_by(ActivityLog.timestamp.desc())
        
        # 分页
        pagination = query.paginate(
            page=page,
            per_page=limit,
            error_out=False
        )
        
        # 构建响应
        logs = [log.to_dict() for log in pagination.items]
        
        return jsonify({
            'code': 0,
            'message': 'success',
            'data': {
                'logs': logs,
                'total_logs': pagination.total,
                'total_pages': pagination.pages,
                'current_page': pagination.page,
                'per_page': limit
            }
        }), 200
        
    except Exception as e:
        current_app.logger.error(f'Get activity logs error: {str(e)}')
        return jsonify({
            'code': 500,
            'message': 'Failed to get activity logs',
            'data': None
        }), 500

@admin_bp.route('/users', methods=['GET', 'OPTIONS'])
@jwt_required()
@role_required(['admin'])
def get_users():
    print("get_users called, decorators order is jwt_required -> role_required")
    """获取用户列表"""
    try:
        # 分页参数
        page = request.args.get('page', 1, type=int)
        page_size = request.args.get('pageSize', 20, type=int)
        
        # 限制每页最大数量
        page_size = min(page_size, 100)
        
        # 搜索参数
        search = request.args.get('search', '').strip()
        role = request.args.get('role', '').strip()
        
        # 构建查询
        query = User.query
        
        if search:
            query = query.filter(User.username.ilike(f'%{search}%'))
        
        if role and role in ['user', 'power_user', 'admin']:
            query = query.filter_by(role=role)
        
        # 排序
        query = query.order_by(User.created_at.desc())
        
        # 分页
        pagination = query.paginate(
            page=page,
            per_page=page_size,
            error_out=False
        )
        
        # 构建响应
        users = []
        for user in pagination.items:
            user_dict = user.to_dict()
            # 添加统计信息
            user_dict['entry_count'] = user.entries.count()
            users.append(user_dict)
        
        return jsonify({
            'code': 0,
            'message': 'success',
            'data': {
                'users': users,
                'total': pagination.total,
                'page': pagination.page,
                'pageSize': page_size,
                'totalPages': pagination.pages
            }
        }), 200
        
    except Exception as e:
        current_app.logger.error(f'Get users error: {str(e)}')
        return jsonify({
            'code': 500,
            'message': 'Failed to get users',
            'data': None
        }), 500

@admin_bp.route('/users/<int:user_id>/role', methods=['PUT', 'OPTIONS'])
@jwt_required()
@role_required(['admin'])
def update_user_role(user_id):
    """更新用户角色"""
    try:
        admin_id = get_jwt_identity()
        
        if user_id == admin_id:
            return jsonify({
                'code': 400,
                'message': 'Cannot change your own role'
            }), 400
        
        user = User.query.get(user_id)
        
        if not user:
            return jsonify({
                'code': 404,
                'message': 'User not found'
            }), 404
        
        data = request.get_json()
        new_role = data.get('role')
        
        if new_role not in ['user', 'power_user', 'admin']:
            return jsonify({
                'code': 400,
                'message': 'Invalid role'
            }), 400
        
        old_role = user.role
        user.role = new_role
        db.session.commit()
        
        # 记录日志
        log_activity(
            'INFO', 
            f'User role changed: {user.username} from {old_role} to {new_role}',
            user_id=admin_id
        )
        
        return jsonify({
            'code': 0,
            'message': 'User role updated successfully',
            'data': user.to_dict()
        }), 200
        
    except Exception as e:
        current_app.logger.error(f'Update user role error: {str(e)}')
        db.session.rollback()
        return jsonify({
            'code': 500,
            'message': 'Failed to update user role',
            'data': None
        }), 500

@admin_bp.route('/users/<int:user_id>', methods=['DELETE', 'OPTIONS'])
@jwt_required()
@role_required(['admin'])
def delete_user(user_id):
    """删除用户"""
    try:
        admin_id = get_jwt_identity()
        
        if user_id == admin_id:
            return jsonify({
                'code': 400,
                'message': 'Cannot delete your own account'
            }), 400
        
        user = User.query.get(user_id)
        
        if not user:
            return jsonify({
                'code': 404,
                'message': 'User not found'
            }), 404
        
        username = user.username
        db.session.delete(user)
        db.session.commit()
        
        # 记录日志
        log_activity(
            'WARNING',
            f'User deleted: {username}',
            user_id=admin_id
        )
        
        return jsonify({
            'code': 0,
            'message': 'User deleted successfully'
        }), 200
        
    except Exception as e:
        current_app.logger.error(f'Delete user error: {str(e)}')
        db.session.rollback()
        return jsonify({
            'code': 500,
            'message': 'Failed to delete user',
            'data': None
        }), 500

@admin_bp.route('/statistics', methods=['GET', 'OPTIONS'])
@jwt_required()
@role_required(['admin'])
def get_statistics():
    """获取系统统计信息"""
    try:
        from sqlalchemy import func
        from datetime import datetime, timedelta
        
        # 用户统计
        total_users = User.query.count()
        users_by_role = db.session.query(
            User.role, 
            func.count(User.id)
        ).group_by(User.role).all()
        
        # 台账统计
        total_entries = LedgerEntry.query.count()
        
        # 最近7天的活动
        week_ago = datetime.utcnow() - timedelta(days=7)
        recent_entries = LedgerEntry.query.filter(
            LedgerEntry.created_at >= week_ago
        ).count()
        
        # 按省份统计
        entries_by_province = db.session.query(
            LedgerEntry.province,
            func.count(LedgerEntry.id)
        ).group_by(LedgerEntry.province).all()
        
        # 构建响应
        return jsonify({
            'code': 0,
            'message': 'success',
            'data': {
                'users': {
                    'total': total_users,
                    'by_role': dict(users_by_role)
                },
                'entries': {
                    'total': total_entries,
                    'recent_7_days': recent_entries,
                    'by_province': dict(entries_by_province)
                }
            }
        }), 200
        
    except Exception as e:
        current_app.logger.error(f'Get statistics error: {str(e)}')
        return jsonify({
            'code': 500,
            'message': 'Failed to get statistics',
            'data': None
        }), 500

@admin_bp.route('/export/logs')
@jwt_required()
@role_required(['admin'])
def export_logs():
    # 获取导出格式
    export_format = request.args.get('format', 'csv').lower()
    # 查询所有日志
    logs = ActivityLog.query.order_by(ActivityLog.timestamp.desc()).all()
    headers = ['ID', '时间', '级别', '用户', '消息', 'IP地址', '详情']
    rows = [
        [
            log.id,
            log.timestamp.strftime('%Y-%m-%d %H:%M:%S') if log.timestamp else '',
            log.level,
            log.username or '',
            log.message,
            log.ip_address or '',
            log.details or ''
        ] for log in logs
    ]
    if export_format == 'excel':
        wb = Workbook()
        ws = wb.active
        ws.append(headers)
        from openpyxl.styles import Font, Alignment
        for cell in ws[1]:
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal='center', vertical='center')
        for row in rows:
            ws.append(row)
        for col in ws.columns:
            max_length = 0
            col_letter = col[0].column_letter
            for cell in col:
                try:
                    max_length = max(max_length, len(str(cell.value)))
                except:
                    pass
            ws.column_dimensions[col_letter].width = max_length + 2
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        return Response(
            output.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={
                'Content-Disposition': 'attachment; filename=logs_export.xlsx',
                'Content-Type': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            }
        )
    elif export_format == 'word':
        doc = Document()
        doc.add_heading('系统日志导出', 0)
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            run = hdr_cells[i].paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(11)
        for row in rows:
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return Response(
            output.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={
                'Content-Disposition': 'attachment; filename=logs_export.docx',
                'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            }
        )
    else:
        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(headers)
        for row in rows:
            writer.writerow(row)
        output.seek(0)
        csv_content = '\ufeff' + output.getvalue()
        if not rows:
            csv_content = '\ufeff' + ','.join(headers) + '\n'
        return Response(
            csv_content,
            mimetype='text/csv',
            headers={
                'Content-Disposition': 'attachment;filename=logs_export.csv',
                'Content-Type': 'text/csv; charset=utf-8'
            }
        )